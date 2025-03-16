"""
CV Analyzer Module

This module analyzes job descriptions to extract relevant skills and requirements,
and then modifies a CV template to match those requirements.
"""

import re
import os
import logging
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
import spacy
import subprocess
import sys

from job_apply_ai.utils.helpers import ensure_directory_exists, sanitize_filename

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Try to load spaCy model, with fallback
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.warning("SpaCy model not found. Using basic text processing instead.")
    nlp = None

class CVAnalyzer:
    """
    A class to analyze job descriptions and extract relevant skills and requirements.
    """
    
    def __init__(self):
        """
        Initialize the CV analyzer.
        """
        # Load spaCy model
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            # If model not found, download it
            logger.warning("SpaCy model not found. Downloading...")
            subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
            self.nlp = spacy.load("en_core_web_sm")
        
        # Define skill categories
        self.skill_categories = {
            "Programming Languages": [
                "python", "java", "javascript", "c++", "c#", "ruby", "php", "swift", 
                "kotlin", "go", "rust", "typescript", "scala", "perl", "r", "matlab",
                "bash", "shell", "powershell", "sql", "html", "css", "dart"
            ],
            "Frameworks & Libraries": [
                "react", "angular", "vue", "django", "flask", "spring", "express", 
                "node.js", "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy",
                "bootstrap", "jquery", "laravel", "symfony", "rails", "asp.net",
                "flutter", "xamarin", ".net", "dotnet", "core", "entity framework"
            ],
            "Databases": [
                "mysql", "postgresql", "mongodb", "sqlite", "oracle", "sql server", 
                "cassandra", "redis", "elasticsearch", "dynamodb", "mariadb", "neo4j",
                "firebase", "supabase", "cockroachdb", "couchdb", "cosmosdb"
            ],
            "Cloud & DevOps": [
                "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "jenkins",
                "terraform", "ansible", "chef", "puppet", "circleci", "travis", "github actions",
                "gitlab ci", "bitbucket pipelines", "heroku", "netlify", "vercel", "digitalocean",
                "linode", "cloudflare", "akamai", "fastly", "lambda", "ec2", "s3", "rds"
            ],
            "Tools & Platforms": [
                "git", "github", "gitlab", "bitbucket", "jira", "confluence", "trello",
                "slack", "notion", "figma", "sketch", "adobe xd", "photoshop", "illustrator",
                "visual studio", "vs code", "intellij", "pycharm", "eclipse", "android studio",
                "xcode", "postman", "insomnia", "swagger", "sentry", "datadog", "grafana"
            ],
            "Methodologies": [
                "agile", "scrum", "kanban", "waterfall", "lean", "tdd", "bdd", "ci/cd",
                "devops", "devsecops", "gitflow", "trunk-based development", "pair programming",
                "extreme programming", "safe", "prince2", "pmp", "itil", "togaf"
            ],
            "Soft Skills": [
                "communication", "teamwork", "leadership", "problem solving", "critical thinking",
                "time management", "adaptability", "creativity", "emotional intelligence",
                "conflict resolution", "negotiation", "presentation", "public speaking",
                "customer service", "mentoring", "coaching", "decision making"
            ],
            "Languages": [
                "english", "german", "french", "spanish", "italian", "portuguese", "dutch",
                "swedish", "norwegian", "danish", "finnish", "russian", "chinese", "japanese",
                "korean", "arabic", "hindi", "bengali", "urdu", "turkish", "polish", "ukrainian"
            ],
            "Business & Analytics": [
                "excel", "powerpoint", "word", "tableau", "power bi", "looker", "google analytics",
                "seo", "sem", "google ads", "facebook ads", "marketing", "sales", "crm", "erp",
                "salesforce", "hubspot", "zoho", "mailchimp", "google workspace", "office 365",
                "financial analysis", "forecasting", "budgeting", "accounting", "quickbooks", "sap"
            ]
        }
    
    def extract_skills_from_description(self, description):
        """
        Extract relevant skills and requirements from a job description.
        
        Args:
            description (str): Job description text.
            
        Returns:
            tuple: (matched_skills, matched_requirements, matched_categories)
        """
        if not description:
            logger.warning("Empty job description provided")
            return [], [], {}
        
        # Process the text
        doc = self.nlp(description.lower())
        
        # Extract all potential skills (nouns and noun phrases)
        potential_skills = []
        for chunk in doc.noun_chunks:
            potential_skills.append(chunk.text)
        
        # Add single tokens that might be skills
        for token in doc:
            if token.pos_ in ["NOUN", "PROPN"]:
                potential_skills.append(token.text)
        
        # Clean up skills
        cleaned_skills = []
        for skill in potential_skills:
            # Remove punctuation and extra whitespace
            cleaned = re.sub(r'[^\w\s]', '', skill).strip()
            if cleaned and len(cleaned) > 1:  # Avoid single characters
                cleaned_skills.append(cleaned)
        
        # Match skills to categories
        matched_skills = set()
        matched_requirements = []
        matched_categories = {}
        
        # Create a flattened list of all skills for faster lookup
        all_skills_flat = {}
        for category, skills in self.skill_categories.items():
            for skill in skills:
                all_skills_flat[skill] = category
        
        # Match skills
        for skill in cleaned_skills:
            # Check for exact matches
            if skill in all_skills_flat:
                matched_skills.add(skill)
                category = all_skills_flat[skill]
                if category not in matched_categories:
                    matched_categories[category] = []
                if skill not in matched_categories[category]:
                    matched_categories[category].append(skill)
            else:
                # Check for partial matches (e.g., "python experience" should match "python")
                for known_skill, category in all_skills_flat.items():
                    if known_skill in skill:
                        matched_skills.add(known_skill)
                        if category not in matched_categories:
                            matched_categories[category] = []
                        if known_skill not in matched_categories[category]:
                            matched_categories[category].append(known_skill)
        
        # Extract requirements (sentences with modal verbs or requirement keywords)
        requirement_keywords = ["required", "must", "should", "need", "essential", "necessary"]
        for sent in doc.sents:
            has_modal = any(token.pos_ == "AUX" and token.dep_ == "aux" for token in sent)
            has_requirement = any(keyword in sent.text.lower() for keyword in requirement_keywords)
            
            if has_modal or has_requirement:
                matched_requirements.append(sent.text)
        
        # If we didn't find any skills, try a more aggressive approach
        if not matched_skills:
            logger.warning("No skills matched using standard approach, trying aggressive matching")
            # Look for any word in the description that matches our skill list
            for word in doc:
                word_text = word.text.lower()
                if word_text in all_skills_flat:
                    matched_skills.add(word_text)
                    category = all_skills_flat[word_text]
                    if category not in matched_categories:
                        matched_categories[category] = []
                    if word_text not in matched_categories[category]:
                        matched_categories[category].append(word_text)
        
        # If still no skills, add some generic skills based on job title keywords
        if not matched_skills and hasattr(doc, 'user_data') and 'job_title' in doc.user_data:
            job_title = doc.user_data['job_title'].lower()
            logger.warning(f"No skills matched, adding generic skills based on job title: {job_title}")
            
            # Map job title keywords to skill categories
            title_to_skills = {
                "developer": ["Programming Languages", "Frameworks & Libraries"],
                "engineer": ["Programming Languages", "Cloud & DevOps"],
                "data": ["Business & Analytics", "Programming Languages"],
                "analyst": ["Business & Analytics", "Databases"],
                "manager": ["Methodologies", "Soft Skills"],
                "designer": ["Tools & Platforms", "Soft Skills"],
                "marketing": ["Business & Analytics", "Soft Skills"],
                "sales": ["Business & Analytics", "Soft Skills"],
                "support": ["Soft Skills", "Tools & Platforms"],
                "admin": ["Tools & Platforms", "Business & Analytics"]
            }
            
            for keyword, categories in title_to_skills.items():
                if keyword in job_title:
                    for category in categories:
                        if category not in matched_categories:
                            matched_categories[category] = []
                        # Add some generic skills from this category
                        skills_to_add = self.skill_categories[category][:3]  # Add first 3 skills
                        for skill in skills_to_add:
                            if skill not in matched_categories[category]:
                                matched_categories[category].append(skill)
                                matched_skills.add(skill)
        
        # Add at least some soft skills if we have no matches
        if not matched_skills:
            logger.warning("No skills matched, adding generic soft skills")
            category = "Soft Skills"
            matched_categories[category] = self.skill_categories[category][:5]  # Add first 5 soft skills
            matched_skills.update(matched_categories[category])
        
        logger.info(f"Extracted {len(matched_skills)} skills across {len(matched_categories)} categories")
        return list(matched_skills), matched_requirements, matched_categories
    
    def process_job_descriptions(self, jobs_df, desc_col="description", title_col="title"):
        """
        Process job descriptions in a DataFrame to extract skills and requirements.
        
        Args:
            jobs_df (DataFrame): DataFrame containing job listings.
            desc_col (str): Column name for job descriptions.
            title_col (str): Column name for job titles.
            
        Returns:
            DataFrame: Updated DataFrame with extracted skills and requirements.
        """
        if jobs_df.empty:
            logger.warning("Empty DataFrame provided")
            return jobs_df
        
        skills_list = []
        requirements_list = []
        categories_list = []
        
        for idx, row in jobs_df.iterrows():
            description_text = str(row.get(desc_col, ""))
            job_title = str(row.get(title_col, "No Title Provided"))
            
            if not description_text.strip():
                logger.warning(f"Empty description for job: {job_title}")
                skills_list.append([])
                requirements_list.append([])
                categories_list.append({})
                continue
            
            matched_skills, matched_requirements, matched_categories = self.extract_skills_from_description(description_text)
            skills_list.append(matched_skills)
            requirements_list.append(matched_requirements)
            categories_list.append(matched_categories)
            
            logger.info(f"Processed job: {job_title} - Found {len(matched_skills)} matching skills")
        
        jobs_df["Extracted Skills"] = skills_list
        jobs_df["Extracted Requirements"] = requirements_list
        jobs_df["Skill Categories"] = categories_list
        
        return jobs_df


class CVModifier:
    """
    A class to modify a CV template based on extracted job requirements.
    """
    
    def __init__(self, cv_template_path):
        """
        Initialize the CV modifier with a template document.
        
        Args:
            cv_template_path (str): Path to the CV template document (.docx).
        """
        self.cv_template_path = cv_template_path
        self.doc = None
        self.load_template()
    
    def load_template(self):
        """
        Load the CV template document.
        """
        try:
            self.doc = Document(self.cv_template_path)
            logger.info(f"Loaded CV template from {self.cv_template_path}")
        except Exception as e:
            logger.error(f"Error loading CV template: {str(e)}")
            raise
    
    def find_skills_section(self):
        """
        Find the skills section in the CV template.
        
        Returns:
            tuple: (paragraph index, paragraph) or (None, None) if not found
        """
        # Common section titles that might indicate skills
        skill_section_keywords = [
            'skills', 'technical skills', 'core skills', 'key skills', 
            'competencies', 'expertise', 'qualifications', 'proficiencies',
            'abilities', 'capabilities', 'technical competencies', 'professional skills',
            'skill set', 'technical expertise', 'core competencies'
        ]
        
        # First try to find an exact heading match
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.lower().strip()
            if text in skill_section_keywords or any(text.startswith(kw) for kw in skill_section_keywords):
                logger.info(f"Found skills section at paragraph {i}: '{para.text}'")
                return i, para
        
        # If no exact match, try to find a paragraph containing skills keywords
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.lower().strip()
            if any(kw in text for kw in skill_section_keywords):
                logger.info(f"Found potential skills section at paragraph {i}: '{para.text}'")
                return i, para
                
        # If still not found, look for bullet points that might contain skill-related words
        skill_related_words = ['proficient', 'experienced', 'knowledge', 'familiar', 'expert', 'advanced', 'intermediate', 'beginner']
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.lower().strip()
            if text.startswith('•') or text.startswith('-') or text.startswith('*'):
                if any(word in text for word in skill_related_words):
                    # This might be part of a skills section, look for a heading above it
                    if i > 0:
                        logger.info(f"Found potential skills bullet point at paragraph {i}: '{para.text}'")
                        # Return the paragraph before this one as it might be the heading
                        return i-1, self.doc.paragraphs[i-1]
        
        logger.warning("Could not find skills section in CV template")
        return None, None
    
    def update_skills_section(self, matched_categories):
        """
        Update the skills section in the CV with matched skills.
        
        Args:
            matched_categories (dict): Dictionary of skill categories and their skills
            
        Returns:
            bool: True if successful, False otherwise
        """
        if not matched_categories:
            logger.warning("No matched categories provided")
            return False
        
        # Find the skills section
        skills_idx, skills_para = self.find_skills_section()
        
        if skills_idx is None:
            # Create a new skills section if one doesn't exist
            logger.info("Creating new skills section in CV")
            skills_para = self.doc.add_paragraph()
            skills_para.style = 'Heading 2'
            skills_para.text = "Skills"
            skills_idx = len(self.doc.paragraphs) - 1
        
        # Clear existing content after the skills heading
        # Find where the next section starts
        next_section_idx = None
        for i in range(skills_idx + 1, len(self.doc.paragraphs)):
            if self.doc.paragraphs[i].style.name.startswith('Heading'):
                next_section_idx = i
                break
        
        # If no next section found, we'll add skills at the end
        if next_section_idx is None:
            next_section_idx = len(self.doc.paragraphs)
        
        # Remove paragraphs between skills heading and next section
        # We need to be careful here as removing paragraphs changes indices
        # So we'll work backwards
        paragraphs_to_remove = list(range(skills_idx + 1, next_section_idx))
        for i in reversed(paragraphs_to_remove):
            if i < len(self.doc.paragraphs):
                p = self.doc.paragraphs[i]
                p._element.getparent().remove(p._element)
        
        # Add matched skills by category
        for category, skills in matched_categories.items():
            if skills:  # Only add categories with skills
                # Add category as subheading
                category_para = self.doc.add_paragraph()
                category_para.style = 'Heading 3'
                category_para.text = category
                
                # Add skills as bullet points
                for skill in skills:
                    skill_para = self.doc.add_paragraph()
                    skill_para.style = 'List Bullet'
                    skill_para.text = skill
        
        logger.info("Successfully updated skills section in CV")
        return True
    
    def save_modified_cv(self, output_path):
        """
        Save the modified CV to the specified path.
        
        Args:
            output_path (str): Path to save the modified CV
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save the document
            self.doc.save(output_path)
            logger.info(f"Saved modified CV to {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving modified CV: {str(e)}")
            return False
    
    def process_multiple_jobs(self, jobs_df, output_dir=None):
        """
        Process multiple jobs and create tailored CVs for each.
        
        Args:
            jobs_df (DataFrame): DataFrame containing job listings with extracted skills.
            output_dir (str, optional): Directory to save the tailored CVs. If None, uses default.
            
        Returns:
            list: List of paths to the generated CV files.
        """
        if jobs_df.empty:
            logger.warning("Empty DataFrame provided")
            return []
        
        if output_dir is None:
            output_dir = os.path.join(os.getcwd(), "job_apply_ai", "outputs", "cvs")
        
        # Ensure the output directory exists
        ensure_directory_exists(output_dir)
        
        # Get current date for filename
        today_date = datetime.today().strftime("%Y-%m-%d")
        
        generated_cvs = []
        
        for idx, row in jobs_df.iterrows():
            job_title = row.get('title', f"Job_{idx+1}")
            company = row.get('company', "Company")
            matched_categories = row.get('Skill Categories', {})
            
            if not matched_categories:
                logger.warning(f"No skills found for job: {job_title} at {company}")
                continue
            
            # Create a fresh copy of the template for each job
            self.doc = Document(self.cv_template_path)
            
            # Update the skills section
            if self.update_skills_section(matched_categories):
                # Generate a filename with date, company, and job title
                safe_company = sanitize_filename(company)
                safe_title = sanitize_filename(job_title)
                filename = f"CV_{today_date}_{safe_company}_{safe_title}.docx"
                output_path = os.path.join(output_dir, filename)
                
                # Save the modified CV
                if self.save_modified_cv(output_path):
                    generated_cvs.append(output_path)
                    logger.info(f"Generated CV for {job_title} at {company}")
            else:
                logger.warning(f"Failed to update CV for job: {job_title} at {company}")
        
        return generated_cvs


def batch_process_jobs(jobs_file, cv_template, output_dir=None):
    """
    Batch process multiple jobs from an Excel file and generate tailored CVs.
    
    Args:
        jobs_file (str): Path to Excel file containing job listings.
        cv_template (str): Path to CV template (.docx).
        output_dir (str, optional): Directory to save the tailored CVs.
        
    Returns:
        list: List of paths to the generated CV files.
    """
    try:
        # Load jobs
        jobs_df = pd.read_excel(jobs_file)
        
        if jobs_df.empty:
            logger.warning(f"No jobs found in {jobs_file}")
            return []
        
        # Process job descriptions
        analyzer = CVAnalyzer()
        processed_df = analyzer.process_job_descriptions(jobs_df)
        
        # Create tailored CVs
        modifier = CVModifier(cv_template)
        generated_cvs = modifier.process_multiple_jobs(processed_df, output_dir)
        
        return generated_cvs
        
    except Exception as e:
        logger.error(f"Error in batch processing: {str(e)}")
        return []


def main():
    """
    Main function to demonstrate the CV analyzer and modifier.
    """
    # Example usage
    import os
    
    # 1. Load job descriptions from Excel
    jobs_file = input("Enter path to jobs Excel file: ")
    if not os.path.exists(jobs_file):
        print(f"File not found: {jobs_file}")
        return
    
    # 2. Load CV template
    cv_template = input("Enter path to your CV template (.docx): ")
    if not os.path.exists(cv_template):
        print(f"File not found: {cv_template}")
        return
    
    # 3. Set output directory
    output_dir = os.path.join(os.getcwd(), "job_apply_ai", "outputs", "cvs")
    
    # 4. Process all jobs and generate CVs
    generated_cvs = batch_process_jobs(jobs_file, cv_template, output_dir)
    
    if generated_cvs:
        print(f"\n✅ Generated {len(generated_cvs)} tailored CVs:")
        for cv_path in generated_cvs:
            print(f"  - {cv_path}")
    else:
        print("\n❌ Failed to generate any CVs")


if __name__ == "__main__":
    main() 