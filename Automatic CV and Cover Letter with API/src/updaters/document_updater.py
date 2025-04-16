"""
Document updater module for tailoring CV and cover letter based on job descriptions.
This module combines document parsing and OpenAI integration to create tailored documents.
"""

import os
import json
import docx
from docx import Document
from typing import Dict, List, Tuple, Any, Optional
import sys
import re
from datetime import datetime

# Add the project root to the path to import our modules
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from parsers.document_parser import CVParser, CoverLetterParser
from utils.openai_integration import OpenAIIntegration


class DocumentUpdater:
    """Class for updating CV and cover letter documents based on job descriptions."""
    
    def __init__(self, cv_path: str, cover_letter_path: str, openai_integration: OpenAIIntegration):
        """Initialize document updater.
        
        Args:
            cv_path: Path to the CV document
            cover_letter_path: Path to the cover letter document
            openai_integration: Initialized OpenAIIntegration instance
        """
        if not os.path.exists(cv_path):
            raise FileNotFoundError(f"CV file not found: {cv_path}")
        if not os.path.exists(cover_letter_path):
            raise FileNotFoundError(f"Cover letter file not found: {cover_letter_path}")
            
        self.cv_parser = CVParser(cv_path)
        self.cover_letter_parser = CoverLetterParser(cover_letter_path)
        self.openai_integration = openai_integration
        self.cv_path = cv_path
        self.cover_letter_path = cover_letter_path
        
    def analyze_job_description(self, job_description: str) -> Dict[str, Any]:
        """Analyze job description and get tailoring suggestions.
        
        Args:
            job_description: Text of the job description
            
        Returns:
            Dictionary with analysis results and suggestions
        """
        # Get current CV content
        cv_content = self.cv_parser.get_all_text()
        
        # Use OpenAI to analyze job description and suggest modifications
        analysis_result = self.openai_integration.analyze_job_description(job_description, cv_content)
        
        return analysis_result
    
    def update_cv(self, job_description: str, output_path: str) -> str:
        """Create a tailored CV based on job description.
        
        Args:
            job_description: Text of the job description
            output_path: Path where the tailored CV should be saved
            
        Returns:
            Path to the tailored CV document
        """
        try:
            # Analyze job description
            analysis = self.analyze_job_description(job_description)
            
            if "error" in analysis:
                raise ValueError(f"Error analyzing job description: {analysis['error']}")
            
            # Parse the raw response if it exists
            suggestions = {}
            if "raw_response" in analysis:
                try:
                    # Try to parse as JSON
                    suggestions = json.loads(analysis["raw_response"])
                except json.JSONDecodeError:
                    # If not valid JSON, extract using regex
                    suggestions = self._extract_suggestions_from_text(analysis["raw_response"])
            else:
                suggestions = analysis  # Already parsed JSON
            
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Create a new document by copying the original
            new_doc = Document(self.cv_path)
            
            # Update profile summary if suggested
            if "profile_summary" in suggestions and suggestions["profile_summary"]:
                self._update_profile_summary(new_doc, suggestions["profile_summary"])
            
            # Update skills if suggested
            if "skills" in suggestions and suggestions["skills"]:
                self._update_skills(new_doc, suggestions["skills"])
            
            # Update experience highlights if suggested
            if "experience_highlights" in suggestions and suggestions["experience_highlights"]:
                self._update_experience(new_doc, suggestions["experience_highlights"])
            
            # Save the updated document
            new_doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error updating CV: {str(e)}")
    
    def update_cover_letter(self, job_description: str, output_path: str) -> str:
        """Create a tailored cover letter based on job description.
        
        Args:
            job_description: Text of the job description
            output_path: Path where the tailored cover letter should be saved
            
        Returns:
            Path to the tailored cover letter document
        """
        try:
            # Get current cover letter and CV content
            cover_letter_content = self.cover_letter_parser.get_all_text()
            cv_content = self.cv_parser.get_all_text()
            
            # Use OpenAI to generate tailored cover letter body
            tailored_body = self.openai_integration.tailor_cover_letter(
                job_description, cover_letter_content, cv_content
            )
            
            if tailored_body.startswith("Error generating cover letter:"):
                raise ValueError(tailored_body)
            
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Create a new document by copying the original
            new_doc = Document(self.cover_letter_path)
            
            # Update the body of the cover letter
            self._update_cover_letter_body(new_doc, tailored_body)
            
            # Save the updated document
            new_doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Error updating cover letter: {str(e)}")
    
    def _extract_suggestions_from_text(self, text: str) -> Dict[str, Any]:
        """Extract suggestions from text when JSON parsing fails.
        
        Args:
            text: Raw text response from OpenAI
            
        Returns:
            Dictionary with extracted suggestions
        """
        suggestions = {
            "profile_summary": "",
            "skills": [],
            "experience_highlights": [],
            "keywords_to_emphasize": []
        }
        
        # Extract profile summary
        profile_match = re.search(r'"profile_summary":\s*"([^"]+)"', text)
        if profile_match:
            suggestions["profile_summary"] = profile_match.group(1)
        
        # Extract skills
        skills_match = re.search(r'"skills":\s*\[(.*?)\]', text, re.DOTALL)
        if skills_match:
            skills_text = skills_match.group(1)
            skills = re.findall(r'"([^"]+)"', skills_text)
            suggestions["skills"] = skills
        
        # Extract experience highlights
        exp_match = re.search(r'"experience_highlights":\s*\[(.*?)\]', text, re.DOTALL)
        if exp_match:
            exp_text = exp_match.group(1)
            experiences = re.findall(r'"([^"]+)"', exp_text)
            suggestions["experience_highlights"] = experiences
        
        # Extract keywords
        keywords_match = re.search(r'"keywords_to_emphasize":\s*\[(.*?)\]', text, re.DOTALL)
        if keywords_match:
            keywords_text = keywords_match.group(1)
            keywords = re.findall(r'"([^"]+)"', keywords_text)
            suggestions["keywords_to_emphasize"] = keywords
        
        return suggestions
    
    def _update_profile_summary(self, doc: Document, new_summary: str) -> None:
        """Update the profile summary in the document.
        
        Args:
            doc: Document object to update
            new_summary: New profile summary text
        """
        # Find the PROFILE section
        profile_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "PROFILE":
                profile_index = i
                break
        
        if profile_index is not None:
            # Clear existing profile paragraphs
            next_section_index = None
            for i in range(profile_index + 1, len(doc.paragraphs)):
                if doc.paragraphs[i].text.strip().isupper() and len(doc.paragraphs[i].text.strip()) < 30:
                    next_section_index = i
                    break
            
            if next_section_index is None:
                next_section_index = len(doc.paragraphs)
            
            # Update the first paragraph after PROFILE
            if profile_index + 1 < len(doc.paragraphs):
                # Clear existing text
                for i in range(profile_index + 1, next_section_index):
                    doc.paragraphs[i].clear()
                
                # Add new summary to the first paragraph after PROFILE
                doc.paragraphs[profile_index + 1].add_run(new_summary)
    
    def _update_skills(self, doc: Document, new_skills: List[str]) -> None:
        """Update the skills section in the document.
        
        Args:
            doc: Document object to update
            new_skills: List of new skills to include
        """
        # Find the SKILLS section (could be SKILLS or TECHNICAL SKILLS)
        skills_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() in ["SKILLS", "TECHNICAL SKILLS"]:
                skills_index = i
                break
        
        if skills_index is not None:
            # Clear existing skills paragraphs
            next_section_index = None
            for i in range(skills_index + 1, len(doc.paragraphs)):
                if doc.paragraphs[i].text.strip().isupper() and len(doc.paragraphs[i].text.strip()) < 30:
                    next_section_index = i
                    break
            
            if next_section_index is None:
                next_section_index = len(doc.paragraphs)
            
            # Update the first paragraph after SKILLS
            if skills_index + 1 < len(doc.paragraphs):
                # Clear existing text
                for i in range(skills_index + 1, next_section_index):
                    doc.paragraphs[i].clear()
                
                # Add new skills to the first paragraph after SKILLS
                skills_text = ", ".join(new_skills)
                doc.paragraphs[skills_index + 1].add_run(skills_text)
    
    def _update_experience(self, doc: Document, new_highlights: List[str]) -> None:
        """Update the experience section in the document to emphasize certain points.
        
        Args:
            doc: Document object to update
            new_highlights: List of experience highlights to emphasize
        """
        # This is a simplified implementation
        # In a real implementation, you would need to carefully modify the experience section
        # while preserving formatting and structure
        
        # Find the EXPERIENCE section (could be EXPERIENCE or WORK EXPERIENCE)
        experience_index = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() in ["EXPERIENCE", "WORK EXPERIENCE"]:
                experience_index = i
                break
        
        if experience_index is not None:
            # Add a note about the highlights at the beginning of the section
            if experience_index + 1 < len(doc.paragraphs):
                highlight_note = "Key highlights relevant to this position: " + "; ".join(new_highlights)
                
                # Insert a new paragraph for the highlights
                p = doc.paragraphs[experience_index]
                run = p.add_run()
                run.add_break()
                run.add_text(highlight_note)
    
    def _update_cover_letter_body(self, doc: Document, new_body: str) -> None:
        """Update the body of the cover letter with new content while preserving formatting.
        
        Args:
            doc: Document object to update
            new_body: New body text for the cover letter
        """
        # Find the start of the body (after greeting) and end (before closing)
        body_start = None
        body_end = None
        
        # Store paragraph formatting for reuse
        format_info = []
        
        # First pass: identify body section and store formatting
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Store formatting information for each paragraph
            format_info.append({
                'style': para.style,
                'alignment': para.alignment,
                'runs': [(run.bold, run.italic, run.underline, run.font.name, run.font.size) 
                        for run in para.runs]
            })
            
            # Skip empty paragraphs but preserve them
            if not text:
                continue
                
            # Look for greeting more carefully
            if any(text.startswith(greeting) for greeting in [
                "Dear ", "To ", "Hi ", "Hello ", "Dear Sir", "Dear Madam", 
                "Dear Hiring", "Dear Recruitment", "Dear HR"
            ]):
                body_start = i + 1
                
            # Look for closing more carefully
            elif any(text.lower().startswith(closing.lower()) for closing in [
                "Sincerely", "Best regards", "Kind regards", "Yours sincerely",
                "Best", "Regards", "Thank you", "Yours faithfully", "Yours truly"
            ]):
                body_end = i
                break
        
        # If we couldn't find the body section, try to make a best guess
        if body_start is None:
            # Look for the first non-header paragraph
            for i, para in enumerate(doc.paragraphs):
                if not any(para.text.strip().lower().startswith(header) for header in [
                    "name:", "address:", "phone:", "email:", "date:"
                ]):
                    body_start = i
                    break
        
        if body_start is None:
            body_start = 0
            
        if body_end is None:
            # Look for signature block
            for i in range(len(doc.paragraphs) - 1, -1, -1):
                text = doc.paragraphs[i].text.strip().lower()
                if text and not any(text.startswith(sig) for sig in [
                    "phone", "email", "address", "mobile", "tel", "website"
                ]):
                    body_end = i
                    break
        
        if body_end is None or body_end <= body_start:
            body_end = len(doc.paragraphs) - 1
        
        # Split new body text into paragraphs
        new_paragraphs = [p.strip() for p in new_body.strip().split("\n\n") if p.strip()]
        
        # Store paragraphs that come before and after the body
        before_body = [para.text for para in doc.paragraphs[:body_start]]
        after_body = [para.text for para in doc.paragraphs[body_end:]]
        
        # Clear the document
        for _ in range(len(doc.paragraphs)):
            if len(doc.paragraphs) > 0:  # Check if there are any paragraphs left
                p = doc.paragraphs[0]._element
                p.getparent().remove(p)
        
        # Rebuild the document
        # Add paragraphs before body
        for i, text in enumerate(before_body):
            p = doc.add_paragraph(text)
            if i < len(format_info):
                self._apply_paragraph_format(p, format_info[i])
        
        # Add new body paragraphs
        for text in new_paragraphs:
            p = doc.add_paragraph(text)
            # Apply a default professional format
            p.style = 'Normal'
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            
        # Add paragraphs after body
        for i, text in enumerate(after_body):
            p = doc.add_paragraph(text)
            if body_end + i < len(format_info):
                self._apply_paragraph_format(p, format_info[body_end + i])
    
    def _apply_paragraph_format(self, paragraph, format_info):
        """Apply stored formatting to a paragraph.
        
        Args:
            paragraph: The paragraph to format
            format_info: Dictionary containing formatting information
        """
        try:
            paragraph.style = format_info['style']
            paragraph.alignment = format_info['alignment']
            
            # If there's text in the paragraph, apply run formatting
            if paragraph.runs and format_info['runs']:
                for run, (bold, italic, underline, font_name, font_size) in zip(
                    paragraph.runs, format_info['runs']
                ):
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = font_size
        except Exception:
            # If any formatting fails, keep going with what we can apply
            pass
