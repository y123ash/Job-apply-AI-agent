"""
Document parser module for extracting content from Word documents.
This module provides functionality to parse CV and cover letter documents.
"""

import docx
from docx import Document
import os
import re
from typing import Dict, List, Tuple, Any


class DocumentParser:
    """Base class for document parsing."""
    
    def __init__(self, file_path: str):
        """Initialize with document file path.
        
        Args:
            file_path: Path to the Word document
        """
        self.file_path = file_path
        self.document = Document(file_path)
        
    def get_all_text(self) -> str:
        """Extract all text from the document.
        
        Returns:
            String containing all text from the document
        """
        return "\n".join([para.text for para in self.document.paragraphs if para.text.strip()])
    
    def save_document(self, output_path: str) -> None:
        """Save the document to a new file.
        
        Args:
            output_path: Path where the document should be saved
        """
        self.document.save(output_path)


class CVParser(DocumentParser):
    """Parser specifically for CV/Resume documents."""
    
    def __init__(self, file_path: str):
        """Initialize CV parser.
        
        Args:
            file_path: Path to the CV document
        """
        super().__init__(file_path)
        self.sections = self._extract_sections()
        
    def _extract_sections(self) -> Dict[str, List[str]]:
        """Extract different sections from the CV.
        
        Returns:
            Dictionary with section names as keys and content as values
        """
        sections = {}
        current_section = "HEADER"
        sections[current_section] = []
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check if this is a section header (all caps, short text)
            if text.isupper() and len(text) < 30 and text not in ["HEADER"]:
                current_section = text
                sections[current_section] = []
            else:
                sections[current_section].append(text)
                
        return sections
    
    def get_personal_info(self) -> Dict[str, str]:
        """Extract personal information from the CV.
        
        Returns:
            Dictionary containing personal information
        """
        personal_info = {}
        
        if "HEADER" in self.sections and self.sections["HEADER"]:
            # First line is typically the name
            if len(self.sections["HEADER"]) > 0:
                personal_info["name"] = self.sections["HEADER"][0]
            
            # Second line typically contains contact information
            if len(self.sections["HEADER"]) > 1:
                contact_line = self.sections["HEADER"][1]
                
                # Extract email using regex
                email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', contact_line)
                if email_match:
                    personal_info["email"] = email_match.group(0)
                
                # Extract phone number
                phone_match = re.search(r'[\+\d][\d\s\(\)-]{8,}', contact_line)
                if phone_match:
                    personal_info["phone"] = phone_match.group(0)
                
                # Extract LinkedIn
                linkedin_match = re.search(r'linkedin\.com/\S+', contact_line)
                if linkedin_match:
                    personal_info["linkedin"] = linkedin_match.group(0)
                    
        return personal_info
    
    def get_profile_summary(self) -> str:
        """Extract profile summary from the CV.
        
        Returns:
            String containing the profile summary
        """
        if "PROFILE" in self.sections:
            return " ".join(self.sections["PROFILE"])
        return ""
    
    def get_education(self) -> List[str]:
        """Extract education information.
        
        Returns:
            List of education entries
        """
        if "EDUCATION" in self.sections:
            return self.sections["EDUCATION"]
        return []
    
    def get_experience(self) -> List[str]:
        """Extract work experience information.
        
        Returns:
            List of experience entries
        """
        if "EXPERIENCE" in self.sections:
            return self.sections["EXPERIENCE"]
        elif "WORK EXPERIENCE" in self.sections:
            return self.sections["WORK EXPERIENCE"]
        return []
    
    def get_skills(self) -> List[str]:
        """Extract skills information.
        
        Returns:
            List of skills entries
        """
        if "SKILLS" in self.sections:
            return self.sections["SKILLS"]
        elif "TECHNICAL SKILLS" in self.sections:
            return self.sections["TECHNICAL SKILLS"]
        return []
    
    def get_all_sections(self) -> Dict[str, List[str]]:
        """Get all extracted sections.
        
        Returns:
            Dictionary with all sections
        """
        return self.sections
    
    def update_section(self, section_name: str, new_content: List[str]) -> None:
        """Update a specific section with new content.
        
        Args:
            section_name: Name of the section to update
            new_content: New content for the section
        """
        if section_name in self.sections:
            # Find the paragraphs that belong to this section
            section_start = None
            section_end = None
            
            for i, para in enumerate(self.document.paragraphs):
                if para.text.strip() == section_name:
                    section_start = i
                    break
            
            if section_start is not None:
                # Find the end of this section (next section header or end of document)
                for i in range(section_start + 1, len(self.document.paragraphs)):
                    if self.document.paragraphs[i].text.strip().isupper() and len(self.document.paragraphs[i].text.strip()) < 30:
                        section_end = i
                        break
                
                if section_end is None:
                    section_end = len(self.document.paragraphs)
                
                # Clear existing paragraphs in this section
                for i in range(section_start + 1, section_end):
                    if self.document.paragraphs[i].text.strip():
                        self.document.paragraphs[i].clear()
                
                # Add new content
                # This is a simplified approach - in a real implementation, you'd need to handle
                # formatting, styles, and potentially add new paragraphs if needed
                current_para = section_start + 1
                for content in new_content:
                    if current_para < section_end:
                        self.document.paragraphs[current_para].add_run(content)
                        current_para += 1
                    else:
                        # Would need to insert new paragraphs here
                        # This is simplified for demonstration
                        pass


class CoverLetterParser(DocumentParser):
    """Parser specifically for cover letter documents."""
    
    def __init__(self, file_path: str):
        """Initialize cover letter parser.
        
        Args:
            file_path: Path to the cover letter document
        """
        super().__init__(file_path)
        self.sections = self._extract_sections()
    
    def _extract_sections(self) -> Dict[str, str]:
        """Extract different parts of the cover letter.
        
        Returns:
            Dictionary with section names as keys and content as values
        """
        sections = {
            "header": "",
            "greeting": "",
            "body": "",
            "closing": ""
        }
        
        # Simple approach based on paragraph position
        # In a real implementation, you'd want more sophisticated parsing
        non_empty_paragraphs = [p for p in self.document.paragraphs if p.text.strip()]
        
        if len(non_empty_paragraphs) >= 1:
            sections["header"] = non_empty_paragraphs[0].text
            
        if len(non_empty_paragraphs) >= 2:
            # Assuming the second paragraph might be a greeting
            if "Dear" in non_empty_paragraphs[1].text or "To" in non_empty_paragraphs[1].text:
                sections["greeting"] = non_empty_paragraphs[1].text
                body_start = 2
            else:
                body_start = 1
                
            # Assuming the last paragraph is the closing
            if len(non_empty_paragraphs) > body_start:
                body_paragraphs = non_empty_paragraphs[body_start:-1] if len(non_empty_paragraphs) > body_start + 1 else []
                sections["body"] = "\n".join([p.text for p in body_paragraphs])
                sections["closing"] = non_empty_paragraphs[-1].text
        
        return sections
    
    def get_header(self) -> str:
        """Get the header section of the cover letter.
        
        Returns:
            String containing the header
        """
        return self.sections.get("header", "")
    
    def get_greeting(self) -> str:
        """Get the greeting section of the cover letter.
        
        Returns:
            String containing the greeting
        """
        return self.sections.get("greeting", "")
    
    def get_body(self) -> str:
        """Get the main body of the cover letter.
        
        Returns:
            String containing the body
        """
        return self.sections.get("body", "")
    
    def get_closing(self) -> str:
        """Get the closing section of the cover letter.
        
        Returns:
            String containing the closing
        """
        return self.sections.get("closing", "")
    
    def update_body(self, new_body: str) -> None:
        """Update the body of the cover letter.
        
        Args:
            new_body: New content for the body section
        """
        # Find paragraphs that contain the body
        non_empty_paragraphs = [i for i, p in enumerate(self.document.paragraphs) if p.text.strip()]
        
        if len(non_empty_paragraphs) >= 3:
            # Assuming body starts at the third paragraph
            body_start = non_empty_paragraphs[2]
            
            # Assuming body ends before the last paragraph
            if len(non_empty_paragraphs) > 3:
                body_end = non_empty_paragraphs[-2]
            else:
                body_end = non_empty_paragraphs[-1]
            
            # Clear existing body paragraphs
            for i in range(body_start, body_end + 1):
                self.document.paragraphs[i].clear()
            
            # Add new body content to the first body paragraph
            self.document.paragraphs[body_start].add_run(new_body)
